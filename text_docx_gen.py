import os
import re
from docx import Document
from openai import OpenAI
import params
import shutil

OPENAI_KEY = params.OPENAI_KEY
client = OpenAI(api_key=OPENAI_KEY)

def call_openai(prompt: str, model: str = "gpt-4", temperature: float = 0.2, max_tokens: int = 1500) -> str:
    """
    Calls the OpenAI API with the provided prompt and returns the resulting text.
    This function is used for both removing irrelevant content and inserting placeholders.
    """
    result_text = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert business quote editor that preserves all document formatting."},
                {"role": "user", "content": prompt}
        ]
    )

    return result_text.choices[0].message.content
    
def ensure_str(text):
    """
    Utility function that ensures the input is a string.
    If the input is not a string, it converts it to one.
    """
    if not isinstance(text, str):
        return str(text)
    return text

def remove_irrelevant_info(tmp_path: str, context: dict) -> None:
    """
    Opens the DOCX document from tmp_path and processes every paragraph and table cell.
    For each text snippet that might contain product details, call a cloud-based LLM
    instructing it to remove irrelevant product details based on the provided context.
    The document is edited in place so that any original line that isn't removed remains unchanged.
    """
    doc = Document(tmp_path)
    
    # Process paragraphs.
    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            prompt = (
                "You are given a snippet from a business quote document. Remove content related to products that are not needed "
                "for this quote based on the context provided. Do not remove universally relevant information such as company details, "
                "terms and conditions, or any textual formatting markers. Return the edited text exactly.\n\n"
                f"Text: {text}\n\n"
                f"Context: {context}"
            )
            new_text = call_openai(prompt)
            new_text = ensure_str(new_text)
            # Only update if the response is non-empty.
            if new_text:
                para.text = new_text

    # Process each cell in all tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                if cell_text.strip():
                    prompt = (
                        "You are given a snippet from a business quote document table cell. Remove content related to products that are irrelevant "
                        "for this quote, while preserving essential company information and formatting. Return the edited text exactly.\n\n"
                        f"Text: {cell_text}\n\n"
                        f"Context: {context}"
                    )
                    new_cell_text = call_openai(prompt)
                    new_cell_text = ensure_str(new_cell_text)
                    if new_cell_text:
                        cell.text = new_cell_text

    doc.save(tmp_path)


# def insert_placeholders(tmp_path: str, context: dict) -> None:
#     """
#     Opens the DOCX file from tmp_path and processes every paragraph and table cell.
#     In each text section, calls the LLM to insert standardized placeholders where dynamic content should be updated.
#     The placeholders are of the format: ---PLACEHOLDER: description---.
#     The document is modified in place without altering any formatting.
#     """
#     doc = Document(tmp_path)
    
#     def process_runs(paragraph):
#         for run in paragraph.runs:
#             text = run.text
#             if text.strip():
#                 prompt = (
#                     f"""
#                     You are given a snippet from a business quote document. Insert standardized placeholders in the text where dynamic content
#                     (such as client email, product descriptions, etc.) should be updated, based on the context provided. Do not remove universally 
#                     relevant information such as company details, terms and conditions, or any textual formatting markers. The placeholders must 
#                     be in the exact format ---PLACEHOLDER: description---. Preserve the formatting, punctuation, and overall structure exactly. 
#                     Do not add extra content at all, only replace with placeholders while keeping the other text exactly the same. Do not reply 
#                     with anything but the text and its edits - no exceptions! If no text is provided then say nothing at all.

#                     For example if the text looked like:
#                         'Product title: xyz'
#                     You might return something like:
#                         'Product title: ---PLACEHOLDER: product title---'
#                     Text: {text}
#                     """
#                 )
#                 new_text = call_openai(prompt)
#                 new_text = ensure_str(new_text)
#                 if new_text:
#                     run.text = new_text

#     # Process paragraphs.
#     for para in doc.paragraphs:
#         process_runs(para)
    
#     # Process each cell in all tables.
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 # A cell can contain multiple paragraphs.
#                 for para in cell.paragraphs:
#                     process_runs(para)

#     base, ext = os.path.splitext(tmp_path)
#     new_path = f"{base}_placeholders{ext}"
#     doc.save(new_path)

# def strip_special_chars(text: str) -> str:
#     """Removes all non-alphanumeric characters for safe dict comparison."""
#     return re.sub(r'[^a-zA-Z0-9]', '', text)

def identify_placeholder_locations(tmp_path: str, context: dict) -> list[str]:
    """
    Opens the DOCX file and extracts all text as a single string.
    Sends the entire text to the LLM to identify snippets that need placeholders.
    Returns a list of strings (exact text snippets) to be replaced or edited.
    """
    doc = Document(tmp_path)
    full_text = []

    # Extract all text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract all text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)

    # Need to remove insidencies of tripple quoates

    prompt_text = """
    You are reviewing a complete quote document to decide where placeholders should be inserted.
    The goal is to identify snippets of text that represent client-specific or dynamic content.
    The input will be a list of runs. 
    Consider the context of the whole document to remove stuff and replace it with placeholders 
    but you can use the list structure to structure your output dict.

    Return ONLY a dict of exact snippets (copied from the text) that either:
    - Should be fully replaced with a placeholder.
    - Should have a placeholder inserted within them.
    If no change is needed for some text, do not include it in the output.
    Maintain the exact casing, punctuation, and structure.

    Use the format: 
    Original snippet: Edited snippet with ---PLACEHOLDER: description---
    for each item of the dict but with one important detail. 
    THE KEYS AND ITEMS NEED TO BE WRAPED IN TRIPPLE QUOTATIONS.
    This is because I need the punctuation to stay the same so I can't remove it but I also don't
    want to mess up the format of the output dictionary.
    I can't show eaxtly how I want the output to look in this prompt as that will mess up my f string.
    So potential output will look something like:
    {\"\"\"Product: xyz\"\"\": \"\"\"Product: ---PLACEHOLDER: product---\"\"\", ...}
    But without the back slashes.

    """
    
    prompt = f"""
    {prompt_text}
    Document text:
    {full_text}
    """
    
    placeholder_instructions = call_openai(prompt)
    # print(raw_response)
    # placeholder_instructions = ensure_str(raw_response).splitlines()
    
    return placeholder_instructions  

def apply_placeholder_edits(tmp_path: str, context: dict) -> None:
    """
    Opens the DOCX file, and for each text run, checks if it matches an original snippet.
    If so, replaces it with the edited version while preserving formatting.
    """
    placeholder_instructions = identify_placeholder_locations(tmp_path, None)
    import ast
    print(placeholder_instructions)
    print(type(placeholder_instructions))
    placeholder_instructions = ast.literal_eval(placeholder_instructions)
    print(placeholder_instructions)
    
    print(placeholder_instructions.keys())
    doc = Document(tmp_path)

    def replace_run_text(run):
        original_text = run.text.strip()
        if original_text in placeholder_instructions:
            run.text = placeholder_instructions[original_text]

    for para in doc.paragraphs:
        for run in para.runs:
            replace_run_text(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        replace_run_text(run)

    base, ext = os.path.splitext(tmp_path)
    new_path = f"{base}_placeholders_applied{ext}"
    doc.save(new_path)

def fill_placeholders(tmp_path: str, context: dict) -> None:
    """
    Opens the DOCX file from tmp_path and searches for standardized placeholders of the form:
       {{PLACEHOLDER: description}}
    Replaces each placeholder with its corresponding value from the context.
    Uses regex to perform a safe replacement so that formatting remains intact.
    The keys in the context are expected to match the placeholder description, normalized to lowercase with underscores.
    """
    doc = Document(tmp_path)
    placeholder_pattern = re.compile(r"\{\{PLACEHOLDER:\s*(.*?)\s*\}\}")

    # Process paragraphs.
    for para in doc.paragraphs:
        text = para.text
        placeholders_found = placeholder_pattern.findall(text)
        new_text = text
        for placeholder in placeholders_found:
            # Ensure the placeholder is treated as a string.
            placeholder = ensure_str(placeholder)
            key = placeholder.lower().replace(" ", "_")
            if key in context:
                replacement = ensure_str(context[key])
                new_text = re.sub(r"\{\{PLACEHOLDER:\s*" + re.escape(placeholder) + r"\s*\}\}", replacement, new_text)
        if new_text != text:
            para.text = new_text

    # Process each cell in tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                placeholders_found = placeholder_pattern.findall(cell_text)
                new_cell_text = cell_text
                for placeholder in placeholders_found:
                    placeholder = ensure_str(placeholder)
                    key = placeholder.lower().replace(" ", "_")
                    if key in context:
                        replacement = ensure_str(context[key])
                        new_cell_text = re.sub(r"\{\{PLACEHOLDER:\s*" + re.escape(placeholder) + r"\s*\}\}", replacement, new_cell_text)
                if new_cell_text != cell_text:
                    cell.text = new_cell_text

    base, ext = os.path.splitext(tmp_path)
    new_path = f"{base}_fill{ext}"
    doc.save(new_path)

def save_intermediate(tmp_path: str, step: str) -> None:
    """
    Copies the current document at tmp_path to a new file that includes the step name.
    This lets you inspect the file after each step.
    """
    base, ext = os.path.splitext(tmp_path)
    new_path = f"{base}_{step}{ext}"
    shutil.copy(tmp_path, new_path)
    print(f"Saved intermediate document: {new_path}")

# ------------------------------------------------------------------------------
# LandGraph Agent Class using the updated workflow
# ------------------------------------------------------------------------------
class LandGraphAgent:
    """
    An agent that processes a .docx file (provided as tmp_path) in three steps:
      1. remove_irrelevant_info: Removes product details not needed for the given quote context.
      2. insert_placeholders: Inserts standardized placeholders (e.g. {{PLACEHOLDER: client_email}})
         in the document wherever dynamic content is required.
      3. fill_placeholders: Replaces these placeholders with the actual content from the context.
    The document is edited directly while preserving its original formatting.
    Intermediate document files are saved after each step.
    """
    def __init__(self):
        self.context = {}

    def update_context(self, context: dict):
        # The context should include both control flags and actual replacement values.
        self.context = context

    def run(self, tmp_path: str) -> str:
        # print("Step 1: Removing irrelevant information from document.")
        # remove_irrelevant_info(tmp_path, self.context)
        # save_intermediate(tmp_path, "step1_removed")
        
        print("Step 2: Inserting placeholders in document.")
        apply_placeholder_edits(tmp_path, self.context)
        
        # print("Step 3: Filling placeholders with provided dynamic content.")
        # fill_placeholders(tmp_path, self.context)
        # save_intermediate(tmp_path, "step3_filled")
        
        # print("Document processing complete. Updated document saved at:", tmp_path)
        return tmp_path

# ------------------------------------------------------------------------------
# Example Workflow Execution
# ------------------------------------------------------------------------------
if __name__ == "__main__":
    # Example temporary DOCX file path (ensure this file exists and is accessible)
    tmp_docx_path = "Tech Solutions Quote Template Ltd.docx"  # Replace with your input document path
    
    # Define the context for both placeholder insertion and replacement.
    # Note: Keys for replacement (for example "client_email") must match the placeholder descriptions
    # normalized to lowercase with underscores.
    context = {
        "client_requirements": {
            "requires_product_a": False,
            "requires_product_b": True
        },
        "client_email": "client@example.com",
        "product_description": "Advanced widget with extended warranty.",
        "company_name": "Company XYZ",
        "terms": "Standard terms and conditions apply."
    }
    
    # Instantiate and run the LandGraph Agent.
    agent = LandGraphAgent()
    agent.update_context(context)
    
    final_doc_path = agent.run(tmp_docx_path)
    print("Final document ready for client at:", final_doc_path)