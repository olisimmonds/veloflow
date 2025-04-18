import docx
import json
from openai import OpenAI
from datetime import date
import params

OPENAI_KEY = params.OPENAI_KEY
client = OpenAI(api_key=OPENAI_KEY)

def extract_doc_structure_docx(doc):
    """
    Extract a structured representation of the document.
    Returns a list of document elements with their type, index information, and text.
    Note: This currently only handles paragraphs and tables. Images and other objects remain untouched.
    """
    elements = []
    # Extract paragraphs with index
    for i, p in enumerate(doc.paragraphs):
        elements.append({
            "type": "paragraph",
            "index": i,
            "text": p.text
        })
    # Extract table cells with table, row, and column indices
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                elements.append({
                    "type": "table",
                    "table_index": ti,
                    "row_index": ri,
                    "col_index": ci,
                    "text": cell.text
                })
    return elements

def replace_text_in_paragraph(paragraph, search_text, replacement_text):
    """
    Iterates through the paragraph's runs to find and replace the search_text.
    Replaces only the first occurrence while preserving the existing formatting.
    """
    for run in paragraph.runs:
        if search_text in run.text:
            run.text = run.text.replace(search_text, replacement_text, 1)
            return True
    return False

def replace_text_in_docx(doc, replacements):
    """
    Iterates over replacement suggestions and applies them based on precise location.
    Each replacement object should include a 'location' key that details the document element location,
    along with 'original' and 'replacement' keys.
    This version edits text at the run level, preserving formatting and leaving other objects (like images) intact.
    """
    for rep in replacements:
        loc = rep.get("location", {})
        original = rep.get("original", "")
        new_text = rep.get("replacement", "")
        if loc.get("type") == "paragraph":
            idx = loc.get("index")
            if idx is not None and idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                if not replace_text_in_paragraph(para, original, new_text):
                    print(f"'{original}' not found in paragraph {idx}")
        elif loc.get("type") == "table":
            ti = loc.get("table_index")
            ri = loc.get("row_index")
            ci = loc.get("col_index")
            try:
                cell = doc.tables[ti].rows[ri].cells[ci]
                replaced = False
                for para in cell.paragraphs:
                    if replace_text_in_paragraph(para, original, new_text):
                        replaced = True
                        break
                if not replaced:
                    print(f"'{original}' not found in table cell {loc}")
            except IndexError:
                print(f"Invalid table location: {loc}")
    return doc

def get_replacements_from_gpt_docx(doc_structure, email, company_context, user_context, user_email):
    """
    Uses GPT-4 (via the OpenAI API) to analyze the document structure and determine which fields need updating.
    Returns a JSON list of objects with the keys:
        - 'location' (dict): where the change should be applied
        - 'original' (str): the exact text to update
        - 'replacement' (str): the new text to insert
    """
    # Convert document structure to JSON string
    doc_structure_json = json.dumps(doc_structure, ensure_ascii=False, indent=2)
    
    prompt = f"""
        You are given a quote document represented as a list of document elements. Each element has a type 
        ('paragraph' or 'table') and corresponding index information, along with its text content. 
        Additionally, consider the following context:

        An email from a potential client:
        {email}

        Company context:
        {company_context}

        User provided context:
        {user_context}

        Todays date:
        {date.today()}

        User's email:
        {user_email}

        Any of the above provided context may be missing or not relevant. It is part of your job to determine what is relevant for completing a client's quote.
        The document may be a template or a previously filled-out template. Your task is to analyze the document elements 
        and complete a quote given the provided context. For each field, provide an object with:
        - 'location': an object indicating where the change should be applied. For a paragraph, include {{'type': 'paragraph', 'index': <paragraph_index>}}. 
        For a table cell, include {{'type': 'table', 'table_index': <table_index>, 'row_index': <row_index>, 'col_index': <col_index>}}.
        - 'original': the exact text snippet to be updated in that element.
        - 'replacement': the new text that should replace the original text.

        Maintain the quote's formatting and structure, only making edits where required to complete the quote. 
        Return your answer as a JSON list.

        Document structure:
        {doc_structure_json}
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a professional sales assistant helping generate and update quotes."},
            {"role": "user", "content": prompt}
        ]
    )

    reply = response.choices[0].message.content
    
    # Extract the JSON portion (assuming it's in a code block with ```json markers)
    reply = reply.split("```json")[-1].split("```")[0]
    try:
        replacements = json.loads(reply)
    except Exception as e:
        print("Error parsing GPT-4 response:", e)
        replacements = []
    return replacements