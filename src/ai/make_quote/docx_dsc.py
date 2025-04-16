import docx
import json
from openai import OpenAI
from datetime import date
import params

# Cloud-ready OpenAI client using an OS-independent library.
OPENAI_KEY = params.OPENAI_KEY
client = OpenAI(api_key=OPENAI_KEY)

def extract_doc_structure(doc):
    """
    Extract a rich representation of the document.
    This version handles paragraphs, tables, and marks images.
    Each element is given an 'id' and a type.
    For images, we note the placeholder and location.
    """
    elements = []
    # Extract paragraphs
    for i, p in enumerate(doc.paragraphs):
        elements.append({
            "id": f"p{i}",
            "type": "paragraph",
            "text": p.text,
            "formatting": {
                # Optionally extract style metadata
                "style": p.style.name if p.style else None
            }
        })
    # Extract tables (with cells)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                elements.append({
                    "id": f"t{ti}_r{ri}_c{ci}",
                    "type": "table-cell",
                    "table_index": ti,
                    "row_index": ri,
                    "col_index": ci,
                    "text": cell.text,
                    "formatting": {}  # Could add style metadata if needed
                })
    # (Optionally) Extract images; 
    # Note: python-docx does not provide an easy way to extract images directly,
    # so you might store a marker which indicates that an image was there.
    # For example, a paragraph may be flagged if it contains an image.
    # You can extend this to search for 'inline shapes' or parts in the document.
    # Example (pseudocode):
    # for i, shape in enumerate(doc.inline_shapes):
    #     elements.append({
    #         "id": f"img{i}",
    #         "type": "image",
    #         "description": shape._inline.graphic.graphicData.uri,  # metadata
    #         "placeholder": f"[Image{i}]"
    #     })
    return elements

def get_new_quote_blueprint(doc_structure, email, company_context, user_context, user_email):
    """
    Sends the full document representation and contexts to GPT to generate a new document blueprint.
    The blueprint can include commands to modify, remove, or add sections.
    The expected return is a JSON list of "elements", where each element includes:
      - id (optional): the original element reference (if modifying an existing element)
      - type: e.g. 'paragraph', 'table-cell', 'image' or 'added-section' (for new parts)
      - text: the new text content (if applicable)
      - formatting: recommendations to preserve style (if applicable)
      - action: one of ['keep', 'modify', 'remove', 'insert_before', 'insert_after']
      - (other metadata as needed)
    """
    doc_structure_json = json.dumps(doc_structure, ensure_ascii=False, indent=2)
    prompt = f"""
    You are provided with a full representation of a quote document. The document is represented as a JSON list of elements.
    Each element represents a paragraph, a table-cell, or an image marker with its content and formatting metadata. 
    See this document here:
    {doc_structure_json}

    Additionally, consider the following context:
    - Email from potential client:
    {email}

    - Company Context:
    {company_context}

    - User Provided Context:
    {user_context}

    - Today's date:
    {date.today()}

    - User's Email:
    {user_email}

    The task is to create a revised quote document blueprint that maintains the overall structure (layout, tone, and style) 
    and preserves core visual elements such as images, while allowing drastic changes in content as needed. 
    This includes adding, removing, or modifying content blocks to produce a finished, polished quote that best reflects 
    the client requirements. For each element, indicate an "action" with one of the following values: 
    'keep', 'modify', 'remove', 'insert_before', or 'insert_after'. 
    For modified or added elements, provide the updated text and any formatting suggestions.

    Return the complete blueprint as a JSON list with the following keys for each element:
    - 'id' (if applicable)
    - 'type'
    - 'action'
    - 'text'
    - 'formatting' (optional)
    """
    response = client.chat.completions.create(
        model="gpt-4o",  # Use a proper GPT-4 endpoint available in the cloud
        messages=[
            {"role": "system", "content": "You are a professional sales assistant expert in creating polished quotes."},
            {"role": "user", "content": prompt}
        ]
    )
    reply = response.choices[0].message.content
    # Expecting the GPT response to be a JSON code block. Extract the JSON portion.
    try:
        json_part = reply.split("```json")[-1].split("```")[0]
        blueprint = json.loads(json_part)
    except Exception as e:
        print("Error parsing GPT response:", e)
        blueprint = []
    return blueprint

def build_doc_from_blueprint(blueprint, source_doc):
    """
    Given a blueprint and the original document, construct a new docx document.
    The blueprint elements indicate what to keep, modify, remove or insert.
    This function reassembles the document by:
      - Copying elements from the source where action is 'keep'
      - Replacing text for elements with action 'modify'
      - Omitting elements with action 'remove'
      - Inserting new paragraphs or sections where indicated by 'insert_before/insert_after'
    """
    new_doc = docx.Document()

    # For simplicity, assume blueprint is a list in the order you want them in the final doc.
    for element in blueprint:
        action = element.get("action")
        elem_type = element.get("type")
        text = element.get("text", "")
        formatting = element.get("formatting", {})

        if action == "remove":
            continue  # Skip this element

        if elem_type == "paragraph":
            if action in ["keep", "modify"] or action.startswith("insert"):
                para = new_doc.add_paragraph(text)
                # Optionally, apply style (if specified by GPT and available in the template)
                if formatting.get("style"):
                    try:
                        para.style = formatting.get("style")
                    except Exception as e:
                        # Fallback if the style is not available
                        pass
        elif elem_type == "table-cell":
            # A simple approach: create a new table (this example creates one cell per element).
            # In practice, you might want to reassemble the table structure using table_index, row_index, etc.
            table = new_doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = text
        elif elem_type == "image":
            # For images, you could reinsert using the stored placeholder info.
            # This assumes that you have a way to locate and attach the image from the source_doc or storage.
            # Example (pseudocode):
            # new_doc.add_picture(get_image_path(element["id"]), width=Inches(4))
            new_doc.add_paragraph(f"[IMAGE: {element.get('description', 'image')}]")
        # You can add other types here as needed.
    return new_doc